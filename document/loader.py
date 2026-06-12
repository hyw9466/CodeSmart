"""文档加载与语义感知分块。支持 .md / .txt / .pdf / .docx"""

from __future__ import annotations

import os
from typing import List, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

# 中英文句子边界分隔符，优先在句号/换行处断开
_SEPARATORS = ["\n\n", "\n", "。", "！", "？", ".", "!", "?", "；", ";", " ", ""]


def load_file(file_path: str) -> str:
    """根据文件扩展名选择解析方式，返回纯文本。"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext}，支持 {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if ext in (".md", ".txt"):
        return _load_text(file_path)
    elif ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    return ""


def _load_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def _load_pdf(file_path: str) -> str:
    """使用 PyMuPDF 提取 PDF 全部页面文本。"""
    import fitz  # pymupdf

    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def _load_docx(file_path: str) -> str:
    """使用 python-docx 提取 Word 文档文本。"""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def validate_and_load(file_path: str) -> Tuple[str, bool, str]:
    """验证文档并加载文本。

    返回：(文本内容, 是否有效, 状态消息)
    """
    ext = os.path.splitext(file_path)[1].lower()
    text = load_file(file_path)
    char_count = len(text.strip())

    if char_count == 0:
        if ext == ".pdf":
            return "", False, "PDF可能是扫描版或图片型PDF，请转换为文本型PDF后重试"
        elif ext == ".docx":
            return "", False, "Word文档内容为空，请检查文档是否有文字内容"
        else:
            return "", False, "文档内容为空"

    return text, True, f"成功提取 {char_count} 个字符"


def split_documents(
    text: str,
    filename: str,
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> List[Document]:
    """将文本按语义感知策略切分为 LangChain Document 列表。"""
    splitter = RecursiveCharacterTextSplitter(
        separators=_SEPARATORS,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    chunks = splitter.split_text(text)
    return [
        Document(
            page_content=chunk,
            metadata={"source": filename, "chunk_index": i},
        )
        for i, chunk in enumerate(chunks)
    ]
